"""
Text extraction from various file formats.

Supports: .txt, .md, .pdf, .docx, .xlsx
"""

from __future__ import annotations

from pathlib import Path


def extract_text(file_path: str | Path) -> str:
    """
    Extract plain text from a file based on its extension.

    Returns the extracted text as a UTF-8 string.
    Raises ValueError for unsupported formats or unreadable files.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    extractors = {
        ".txt": _extract_txt,
        ".md": _extract_txt,
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".xlsx": _extract_xlsx,
    }

    extractor = extractors.get(ext)
    if extractor is None:
        raise ValueError(f"Unsupported format: {ext}")

    return extractor(path)


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """
    Extract plain text from file content given as bytes.

    Useful for uploaded files that haven't been saved to disk yet.
    """
    ext = Path(filename).suffix.lower()

    extractors = {
        ".txt": lambda _: content.decode("utf-8"),
        ".md": lambda _: content.decode("utf-8"),
        ".pdf": _extract_pdf_bytes,
        ".docx": _extract_docx_bytes,
        ".xlsx": _extract_xlsx_bytes,
    }

    extractor = extractors.get(ext)
    if extractor is None:
        raise ValueError(f"Unsupported format: {ext}")

    return extractor(content)


# ── Format-specific extractors ──────────────────────────────────────

def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(path))
    except Exception as e:
        raise ValueError(f"Failed to read PDF: {e}")

    pages = []
    for page in reader.pages:
        try:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text)
        except Exception:
            continue

    if not pages:
        raise ValueError("PDF contains no extractable text (may be image-only)")

    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def _extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(path), read_only=True, data_only=True)
    sheets_text = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            # Skip completely empty rows
            if any(c.strip() for c in cells):
                rows.append(" | ".join(cells))
        if rows:
            sheets_text.append(f"## {sheet_name}\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets_text)


def _extract_pdf_bytes(content: bytes) -> str:
    from pypdf import PdfReader
    from io import BytesIO

    try:
        reader = PdfReader(BytesIO(content))
    except Exception as e:
        raise ValueError(f"Failed to read PDF: {e}")

    pages = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text)
        except Exception:
            # Skip pages that can't be extracted
            continue

    if not pages:
        raise ValueError("PDF contains no extractable text (may be image-only)")

    return "\n\n".join(pages)


def _extract_docx_bytes(content: bytes) -> str:
    from docx import Document
    from io import BytesIO
    doc = Document(BytesIO(content))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def _extract_xlsx_bytes(content: bytes) -> str:
    from openpyxl import load_workbook
    from io import BytesIO
    wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
    sheets_text = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                rows.append(" | ".join(cells))
        if rows:
            sheets_text.append(f"## {sheet_name}\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets_text)
